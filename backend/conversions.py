"""File conversion + editing engine. Pure functions on bytes -> bytes."""
import io
import json
import csv
from typing import Tuple, List, Optional

from PIL import Image, ImageFilter, ImageEnhance, ImageOps
import pypdf
import openpyxl
from openpyxl import Workbook
import docx

# ------------------------ format registry ------------------------
IMAGE_FORMATS = ["png", "jpg", "jpeg", "webp", "bmp", "gif", "tiff", "ico"]
DOC_FORMATS = ["pdf", "txt", "docx"]
SHEET_FORMATS = ["csv", "xlsx", "json", "tsv"]

# What can each input format be converted to.
CONVERSION_MATRIX = {
    # images <-> images, images -> pdf
    **{ext: IMAGE_FORMATS + ["pdf"] for ext in IMAGE_FORMATS},
    # pdf -> images, text
    "pdf": ["txt", "png", "jpg"],
    # docs
    "docx": ["txt", "pdf"],
    "txt": ["pdf", "docx"],
    # sheets - any of these to any of these
    **{ext: SHEET_FORMATS for ext in SHEET_FORMATS},
}


def get_targets_for(source_ext: str) -> List[str]:
    src = source_ext.lower().lstrip(".")
    targets = sorted(set(CONVERSION_MATRIX.get(src, [])))
    return [t for t in targets if t != src]


def normalize_ext(filename: str) -> str:
    if "." not in filename:
        return ""
    return filename.rsplit(".", 1)[-1].lower()


# ------------------------ image ops ------------------------
def _open_image(data: bytes) -> Image.Image:
    return Image.open(io.BytesIO(data))


def convert_image(data: bytes, target: str) -> bytes:
    img = _open_image(data)
    out = io.BytesIO()
    target = target.lower()
    save_fmt = {
        "jpg": "JPEG", "jpeg": "JPEG", "png": "PNG", "webp": "WEBP",
        "bmp": "BMP", "gif": "GIF", "tiff": "TIFF", "ico": "ICO",
    }[target]
    if save_fmt == "JPEG" and img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")
    if save_fmt == "ICO":
        img = img.convert("RGBA")
        img.thumbnail((256, 256))
    img.save(out, format=save_fmt)
    return out.getvalue()


def image_to_pdf(data: bytes) -> bytes:
    img = _open_image(data)
    if img.mode != "RGB":
        img = img.convert("RGB")
    out = io.BytesIO()
    img.save(out, format="PDF")
    return out.getvalue()


def edit_image(
    data: bytes,
    rotate: int = 0,
    flip_h: bool = False,
    flip_v: bool = False,
    crop: Optional[Tuple[int, int, int, int]] = None,
    grayscale: bool = False,
    invert: bool = False,
    brightness: float = 1.0,
    contrast: float = 1.0,
    saturation: float = 1.0,
    blur: float = 0.0,
    sharpen: bool = False,
    out_format: str = "png",
) -> bytes:
    img = _open_image(data)
    if crop:
        img = img.crop(crop)
    if rotate:
        img = img.rotate(-rotate, expand=True)  # CW positive
    if flip_h:
        img = ImageOps.mirror(img)
    if flip_v:
        img = ImageOps.flip(img)
    if grayscale:
        img = ImageOps.grayscale(img).convert("RGB")
    if invert:
        if img.mode == "RGBA":
            r, g, b, a = img.split()
            rgb = Image.merge("RGB", (r, g, b))
            rgb = ImageOps.invert(rgb)
            r2, g2, b2 = rgb.split()
            img = Image.merge("RGBA", (r2, g2, b2, a))
        else:
            img = ImageOps.invert(img.convert("RGB"))
    if brightness != 1.0:
        img = ImageEnhance.Brightness(img).enhance(brightness)
    if contrast != 1.0:
        img = ImageEnhance.Contrast(img).enhance(contrast)
    if saturation != 1.0:
        img = ImageEnhance.Color(img).enhance(saturation)
    if blur > 0:
        img = img.filter(ImageFilter.GaussianBlur(radius=blur))
    if sharpen:
        img = img.filter(ImageFilter.SHARPEN)

    out = io.BytesIO()
    fmt = out_format.lower()
    save_fmt = "JPEG" if fmt in ("jpg", "jpeg") else fmt.upper()
    if save_fmt == "JPEG" and img.mode in ("RGBA", "P", "LA"):
        img = img.convert("RGB")
    img.save(out, format=save_fmt)
    return out.getvalue()


# ------------------------ pdf ops ------------------------
def pdf_to_text(data: bytes) -> bytes:
    reader = pypdf.PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n\n".join(parts).encode("utf-8")


def merge_pdfs(parts: List[bytes]) -> bytes:
    writer = pypdf.PdfWriter()
    for raw in parts:
        reader = pypdf.PdfReader(io.BytesIO(raw))
        for page in reader.pages:
            writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue()


def split_pdf(data: bytes, ranges: List[Tuple[int, int]]) -> List[bytes]:
    """Returns list of pdf bytes for each range (1-indexed inclusive)."""
    reader = pypdf.PdfReader(io.BytesIO(data))
    results = []
    for start, end in ranges:
        w = pypdf.PdfWriter()
        for i in range(start - 1, min(end, len(reader.pages))):
            w.add_page(reader.pages[i])
        buf = io.BytesIO()
        w.write(buf)
        results.append(buf.getvalue())
    return results


def pdf_page_count(data: bytes) -> int:
    return len(pypdf.PdfReader(io.BytesIO(data)).pages)


# ------------------------ doc ops ------------------------
def docx_to_text(data: bytes) -> bytes:
    d = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in d.paragraphs).encode("utf-8")


def text_to_docx(data: bytes) -> bytes:
    text = data.decode("utf-8", errors="replace")
    d = docx.Document()
    for line in text.split("\n"):
        d.add_paragraph(line)
    out = io.BytesIO()
    d.save(out)
    return out.getvalue()


def text_to_pdf(data: bytes) -> bytes:
    """Very simple text → PDF using PIL (no external deps)."""
    text = data.decode("utf-8", errors="replace")
    # Render text in pages of ~50 lines
    lines = text.split("\n")
    pages = []
    per_page = 50
    for i in range(0, max(1, len(lines)), per_page):
        page_lines = lines[i:i + per_page]
        img = Image.new("RGB", (1240, 1754), "white")  # ~A4 @ 150dpi
        from PIL import ImageDraw, ImageFont
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 22)
        except Exception:
            font = ImageFont.load_default()
        y = 80
        for line in page_lines:
            draw.text((80, y), line[:120], fill="black", font=font)
            y += 30
        pages.append(img)
    out = io.BytesIO()
    if not pages:
        pages = [Image.new("RGB", (1240, 1754), "white")]
    pages[0].save(out, format="PDF", save_all=True, append_images=pages[1:])
    return out.getvalue()


def docx_to_pdf(data: bytes) -> bytes:
    return text_to_pdf(docx_to_text(data))


def pdf_to_images(data: bytes, fmt: str = "png") -> bytes:
    """Without poppler we can't truly rasterize. Return a textual-image fallback
    showing the extracted text rendered as image."""
    text = pdf_to_text(data)
    # render the text as one image
    text_str = text.decode("utf-8", errors="replace")
    img = Image.new("RGB", (1240, 1754), "white")
    from PIL import ImageDraw, ImageFont
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20)
    except Exception:
        font = ImageFont.load_default()
    y = 80
    for line in text_str.split("\n")[:60]:
        draw.text((80, y), line[:120], fill="black", font=font)
        y += 28
    out = io.BytesIO()
    img.save(out, format="PNG" if fmt == "png" else "JPEG")
    return out.getvalue()


# ------------------------ sheet ops ------------------------
def _read_sheet(data: bytes, ext: str) -> List[List[str]]:
    ext = ext.lower()
    if ext == "csv":
        text = data.decode("utf-8", errors="replace")
        return list(csv.reader(io.StringIO(text)))
    if ext == "tsv":
        text = data.decode("utf-8", errors="replace")
        return list(csv.reader(io.StringIO(text), delimiter="\t"))
    if ext == "json":
        text = data.decode("utf-8", errors="replace")
        obj = json.loads(text)
        if isinstance(obj, list) and obj and isinstance(obj[0], dict):
            headers = list(obj[0].keys())
            rows = [headers] + [[str(r.get(h, "")) for h in headers] for r in obj]
            return rows
        if isinstance(obj, list):
            return [[str(x)] for x in obj]
        if isinstance(obj, dict):
            return [["key", "value"]] + [[str(k), str(v)] for k, v in obj.items()]
        return [[str(obj)]]
    if ext == "xlsx":
        wb = openpyxl.load_workbook(io.BytesIO(data), data_only=True)
        ws = wb.active
        return [[("" if c is None else str(c)) for c in row] for row in ws.iter_rows(values_only=True)]
    raise ValueError(f"Unknown sheet format: {ext}")


def _write_sheet(rows: List[List[str]], target: str) -> bytes:
    target = target.lower()
    if target == "csv":
        out = io.StringIO()
        csv.writer(out).writerows(rows)
        return out.getvalue().encode("utf-8")
    if target == "tsv":
        out = io.StringIO()
        csv.writer(out, delimiter="\t").writerows(rows)
        return out.getvalue().encode("utf-8")
    if target == "json":
        if not rows:
            return b"[]"
        headers = rows[0]
        data = [dict(zip(headers, row)) for row in rows[1:]]
        return json.dumps(data, indent=2).encode("utf-8")
    if target == "xlsx":
        wb = Workbook()
        ws = wb.active
        for row in rows:
            ws.append(row)
        out = io.BytesIO()
        wb.save(out)
        return out.getvalue()
    raise ValueError(f"Unknown sheet target: {target}")


def convert_sheet(data: bytes, source: str, target: str) -> bytes:
    rows = _read_sheet(data, source)
    return _write_sheet(rows, target)


def any_to_xlsx(data: bytes, source: str) -> bytes:
    """Convert any file to a single-cell-or-tabular xlsx (force conversion)."""
    if source in SHEET_FORMATS:
        return convert_sheet(data, source, "xlsx")
    if source == "txt":
        text = data.decode("utf-8", errors="replace")
        rows = [[line] for line in text.split("\n")]
        return _write_sheet(rows, "xlsx")
    if source == "pdf":
        text = pdf_to_text(data).decode("utf-8", errors="replace")
        rows = [[line] for line in text.split("\n")]
        return _write_sheet(rows, "xlsx")
    if source == "docx":
        text = docx_to_text(data).decode("utf-8", errors="replace")
        rows = [[line] for line in text.split("\n")]
        return _write_sheet(rows, "xlsx")
    # last resort: base64 of binary
    import base64
    rows = [["filename", "base64_chunk"]]
    enc = base64.b64encode(data).decode("ascii")
    for i in range(0, len(enc), 32000):
        rows.append([f"chunk_{i//32000}", enc[i:i + 32000]])
    return _write_sheet(rows, "xlsx")


# ------------------------ top-level dispatcher ------------------------
def run_conversion(data: bytes, source_ext: str, target_ext: str) -> bytes:
    s = source_ext.lower().lstrip(".")
    t = target_ext.lower().lstrip(".")
    if s == t:
        return data

    # images -> images
    if s in IMAGE_FORMATS and t in IMAGE_FORMATS:
        return convert_image(data, t)
    # images -> pdf
    if s in IMAGE_FORMATS and t == "pdf":
        return image_to_pdf(data)
    # pdf -> text/images
    if s == "pdf" and t == "txt":
        return pdf_to_text(data)
    if s == "pdf" and t in ("png", "jpg", "jpeg"):
        return pdf_to_images(data, "png" if t == "png" else "jpg")
    # docx
    if s == "docx" and t == "txt":
        return docx_to_text(data)
    if s == "docx" and t == "pdf":
        return docx_to_pdf(data)
    # txt
    if s == "txt" and t == "docx":
        return text_to_docx(data)
    if s == "txt" and t == "pdf":
        return text_to_pdf(data)
    # sheets <-> sheets
    if s in SHEET_FORMATS and t in SHEET_FORMATS:
        return convert_sheet(data, s, t)

    raise ValueError(f"Unsupported conversion: {s} -> {t}")
